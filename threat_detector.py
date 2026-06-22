from docx import Document
from PIL import Image
import pytesseract
import os
import re

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# Extract text from DOCX
def extract_text(doc_path):

    doc = Document(doc_path)

    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    return " ".join(full_text)


# Detect harmful words
def detect_harmful_words(text):

    with open("data/harmful_words.txt", "r", encoding="utf-8") as f:
        harmful_words = f.read().splitlines()

    found = []

    text = text.lower()

    for word in harmful_words:
        if word.lower() in text:
            found.append(word)

    return found


# Extract URLs
def extract_urls(text):

    pattern = r'https?://\S+'

    return re.findall(pattern, text)


# Check suspicious domains
def suspicious_links(urls):

    with open("data/suspicious_domains.txt", "r", encoding="utf-8") as f:
        domains = f.read().splitlines()

    bad_links = []

    for url in urls:
        for domain in domains:
            if domain.lower() in url.lower():
                bad_links.append(url)

    return list(set(bad_links))


# Count images
def count_images(doc_path):

    doc = Document(doc_path)

    image_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    return image_count


# Detect executable file names
def detect_executables(text):

    pattern = r'\b\S+\.(?:exe|bat|vbs|js|cmd)\b'

    matches = re.finditer(pattern, text, re.IGNORECASE)

    return [match.group() for match in matches]


# Scan images separately
def scan_images_for_harmful_text(doc_path):

    doc = Document(doc_path)

    harmful_keywords = [
        "hack",
        "malware",
        "virus",
        "trojan",
        "ransom",
        "password",
        "keylogger",
        "phishing",
        "exploit",
        "spyware",
        "backdoor"
    ]

    image_results = []

    image_number = 1

    for rel in doc.part.rels.values():

        if "image" in rel.target_ref:

            try:

                image_data = rel.target_part.blob

                temp_image = f"temp_image_{image_number}.png"

                with open(temp_image, "wb") as f:
                    f.write(image_data)

                extracted_text = pytesseract.image_to_string(
                    Image.open(temp_image)
                ).lower()

                found_keywords = []

                for word in harmful_keywords:

                    if word in extracted_text:
                        found_keywords.append(word)

                image_results.append({
                    "image_no": image_number,
                    "keywords": found_keywords,
                    "ocr_text": extracted_text.strip()
                })

                if os.path.exists(temp_image):
                    os.remove(temp_image)

                image_number += 1

            except Exception:
                pass

    return image_results


# Calculate risk score
def calculate_risk(harmful, bad_links, images):

    score = 0

    score += len(harmful) * 10
    score += len(bad_links) * 20

    if images > 0:
        score += images * 5

    return score


# Classification
def classify(score):

    if score < 20:
        return "SAFE"

    elif score < 50:
        return "SUSPICIOUS"

    else:
        return "HIGH RISK"